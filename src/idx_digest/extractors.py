from __future__ import annotations

import mimetypes
import re
import tempfile
import time
import warnings
import zipfile
from dataclasses import dataclass
from pathlib import Path, PurePosixPath

import openpyxl
import pymupdf
import pytesseract
from bs4 import BeautifulSoup
from docx import Document
from PIL import Image

from .config import Settings
from .observability import RunObserver


ZIP_MAX_MEMBERS = 100
ZIP_MAX_SUPPORTED_MEMBERS = 50
ZIP_MAX_MEMBER_UNCOMPRESSED_BYTES = 50_000_000
ZIP_MAX_TOTAL_UNCOMPRESSED_BYTES = 100_000_000
ZIP_MAX_COMPRESSION_RATIO = 200.0
ZIP_SUPPORTED_SUFFIXES = {
    ".pdf",
    ".xlsx",
    ".xlsm",
    ".docx",
    ".html",
    ".htm",
    ".txt",
    ".csv",
    ".json",
    ".xml",
}


@dataclass(frozen=True)
class ExtractionResult:
    text: str
    method: str


def normalize_text(text: str) -> str:
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()


def extract_pdf(path: Path, settings: Settings, observer: RunObserver | None = None) -> ExtractionResult:
    doc = pymupdf.open(path)
    parts: list[str] = []
    used_ocr = False
    task_id = observer.start_task(
        f"Extracting PDF pages • {path.name}", total=len(doc), kind="items"
    ) if observer else None
    started = time.perf_counter()
    try:
        for page_no, page in enumerate(doc, start=1):
            page_started = time.perf_counter()
            native = page.get_text("text", sort=True).strip()
            ocr_chars = 0
            if len(native) >= settings.ocr_min_chars_per_page:
                page_text = native
                method = "native"
            else:
                if observer:
                    observer.event(
                        "extract",
                        "native PDF text sparse; running OCR",
                        file=path.name,
                        page=page_no,
                        native_characters=len(native),
                        threshold=settings.ocr_min_chars_per_page,
                    )
                pix = page.get_pixmap(matrix=pymupdf.Matrix(2.0, 2.0), alpha=False)
                image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                ocr = pytesseract.image_to_string(image, lang=settings.ocr_lang).strip()
                ocr_chars = len(ocr)
                page_text = ocr if len(ocr) > len(native) else native
                used_ocr = used_ocr or bool(ocr)
                method = "ocr" if len(ocr) > len(native) else "native"
            parts.append(f"\n\n===== PAGE {page_no} =====\n{page_text}")
            if observer:
                observer.update_task(task_id, advance=1)
                observer.event(
                    "extract",
                    "PDF page extracted",
                    file=path.name,
                    page=page_no,
                    method=method,
                    native_characters=len(native),
                    ocr_characters=ocr_chars,
                    output_characters=len(page_text),
                    elapsed_seconds=f"{time.perf_counter() - page_started:.3f}",
                )
    finally:
        doc.close()
        if observer:
            observer.finish_task(task_id)
    method = "pdf-native+ocr" if used_ocr else "pdf-native"
    result = ExtractionResult(normalize_text("".join(parts)), method)
    if observer:
        observer.event(
            "extract",
            "PDF extraction complete",
            file=path.name,
            pages=len(parts),
            method=method,
            characters=len(result.text),
            elapsed_seconds=f"{time.perf_counter() - started:.3f}",
        )
    return result


def _xlsx_sheet_priority(sheet) -> tuple[int, str]:
    title = str(sheet.title or "").strip()
    title_low = title.lower()
    taxonomy_terms = ("inlinexbrl", "inline xbrl", "token", "hidden", "taxonomy", "schema", "xbrl", "contexts", "units")
    primary_terms = (
        "statement of financial position", "financial position", "balance sheet", "posisi keuangan",
        "income statement", "profit or loss", "profit and loss", "laba rugi", "pendapatan",
        "cash flow", "arus kas", "changes in equity", "perubahan ekuitas",
        "statement of comprehensive income", "penghasilan komprehensif",
    )
    supporting_terms = ("notes", "catatan", "segment", "segmen", "context")
    if any(term in title_low for term in taxonomy_terms):
        return (3, title_low)
    if any(term in title_low for term in primary_terms):
        return (0, title_low)
    if any(term in title_low for term in supporting_terms):
        return (2, title_low)

    # Numeric IDX taxonomy sheet names are common, so inspect only a small sample
    # to determine ordering. Full extraction below still preserves every sheet.
    sample: list[str] = []
    try:
        for row_no, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            for value in row:
                if value is not None and str(value).strip():
                    sample.append(str(value).strip().lower())
            if row_no >= 25 or len(sample) >= 80:
                break
    except Exception:
        pass
    joined = " | ".join(sample)
    if any(term in joined for term in primary_terms):
        return (0, title_low)
    if any(term in joined for term in taxonomy_terms):
        return (3, title_low)
    if any(term in joined for term in supporting_terms):
        return (2, title_low)
    return (1, title_low)


def extract_xlsx(path: Path, settings: Settings, observer: RunObserver | None = None) -> ExtractionResult:
    # Load both cached values and formulas. IDX workbooks occasionally contain
    # formulas without cached values; exposing the formula is safer than silently
    # turning the cell into blank.
    with warnings.catch_warnings():
        warnings.filterwarnings(
            "ignore",
            message="Data Validation extension is not supported and will be removed",
            category=UserWarning,
        )
        workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
        formula_workbook = openpyxl.load_workbook(path, read_only=True, data_only=False)

    output: list[str] = []
    cells_seen = 0
    truncated = False
    formula_fallbacks = 0
    ordered_sheets = sorted(workbook.worksheets, key=_xlsx_sheet_priority)
    formula_by_title = {sheet.title: sheet for sheet in formula_workbook.worksheets}
    task_id = observer.start_task(
        f"Extracting spreadsheet • {path.name}", total=len(ordered_sheets), kind="items"
    ) if observer else None
    started = time.perf_counter()
    try:
        for sheet in ordered_sheets:
            sheet_started = time.perf_counter()
            priority = _xlsx_sheet_priority(sheet)[0]
            priority_label = {0: "PRIMARY", 1: "GENERAL", 2: "SUPPORTING", 3: "TAXONOMY"}.get(priority, "GENERAL")
            output.append(f"\n\n===== SHEET: {sheet.title} =====\nSHEET PRIORITY: {priority_label}")
            rows_seen = 0
            meaningful_rows = 0
            formula_sheet = formula_by_title.get(sheet.title)
            formula_rows = formula_sheet.iter_rows(values_only=False) if formula_sheet is not None else None
            for data_row in sheet.iter_rows(values_only=False):
                rows_seen += 1
                formula_row = next(formula_rows, ()) if formula_rows is not None else ()
                values: list[str] = []
                for index, data_cell in enumerate(data_row):
                    value = data_cell.value
                    formula_value = formula_row[index].value if index < len(formula_row) else None
                    if value is None and isinstance(formula_value, str) and formula_value.startswith("="):
                        value = f"[FORMULA_NO_CACHED_VALUE {data_cell.coordinate}: {formula_value}]"
                        formula_fallbacks += 1
                    values.append("" if value is None else str(value))
                nonempty = [index for index, value in enumerate(values) if value.strip()]
                if not nonempty:
                    continue
                values = values[: nonempty[-1] + 1]
                meaningful_rows += 1
                cells_seen += len(values)
                if cells_seen > settings.max_xlsx_cells:
                    truncated = True
                    break
                output.append(f"ROW {rows_seen} | " + " | ".join(values))
            if observer:
                observer.update_task(task_id, advance=1)
                observer.event(
                    "extract", "spreadsheet sheet extracted", file=path.name,
                    sheet=sheet.title, sheet_priority=priority_label, rows=rows_seen,
                    meaningful_rows=meaningful_rows, cells_seen=cells_seen,
                    formula_fallbacks=formula_fallbacks,
                    elapsed_seconds=f"{time.perf_counter() - sheet_started:.3f}",
                )
            if truncated:
                break
    finally:
        workbook.close()
        formula_workbook.close()
        if observer:
            observer.finish_task(task_id)
    if truncated:
        output.append("\n[TRUNCATED: spreadsheet cell limit reached after priority ordering]")
    result = ExtractionResult(normalize_text("\n".join(output)), "xlsx-priority+formula-fallback")
    if observer:
        observer.event(
            "extract", "spreadsheet extraction complete", file=path.name,
            cells=cells_seen, truncated=truncated, characters=len(result.text),
            formula_fallbacks=formula_fallbacks,
            elapsed_seconds=f"{time.perf_counter() - started:.3f}",
        )
    return result


def extract_docx(path: Path, observer: RunObserver | None = None) -> ExtractionResult:
    started = time.perf_counter()
    document = Document(path)
    chunks = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            chunks.append("\t".join(cell.text for cell in row.cells))
    result = ExtractionResult(normalize_text("\n".join(chunks)), "docx")
    if observer:
        observer.event(
            "extract",
            "DOCX extraction complete",
            file=path.name,
            paragraphs=len(document.paragraphs),
            tables=len(document.tables),
            characters=len(result.text),
            elapsed_seconds=f"{time.perf_counter() - started:.3f}",
        )
    return result


def extract_html(path: Path, observer: RunObserver | None = None) -> ExtractionResult:
    started = time.perf_counter()
    soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="replace"), "html.parser")
    result = ExtractionResult(normalize_text(soup.get_text("\n")), "html")
    if observer:
        observer.event(
            "extract",
            "HTML extraction complete",
            file=path.name,
            characters=len(result.text),
            elapsed_seconds=f"{time.perf_counter() - started:.3f}",
        )
    return result


def extract_text(path: Path, observer: RunObserver | None = None) -> ExtractionResult:
    started = time.perf_counter()
    result = ExtractionResult(
        normalize_text(path.read_text(encoding="utf-8", errors="replace")), "text"
    )
    if observer:
        observer.event(
            "extract",
            "text extraction complete",
            file=path.name,
            characters=len(result.text),
            elapsed_seconds=f"{time.perf_counter() - started:.3f}",
        )
    return result


def _validate_zip_member(info: zipfile.ZipInfo) -> PurePosixPath:
    raw_name = info.filename.replace("\\", "/")
    member_path = PurePosixPath(raw_name)
    if member_path.is_absolute() or ".." in member_path.parts:
        raise ValueError(f"Unsafe ZIP member path: {info.filename!r}")
    unix_mode = (info.external_attr >> 16) & 0o170000
    if unix_mode == 0o120000:
        raise ValueError(f"ZIP symlink member is not allowed: {info.filename!r}")
    if info.flag_bits & 0x1:
        raise ValueError(f"Encrypted ZIP member is not supported: {info.filename!r}")
    if info.file_size < 0 or info.file_size > ZIP_MAX_MEMBER_UNCOMPRESSED_BYTES:
        raise ValueError(f"ZIP member exceeds safe size limit: {info.filename!r}")
    if info.file_size > 0:
        if info.compress_size <= 0:
            raise ValueError(f"ZIP member has suspicious compressed size: {info.filename!r}")
        ratio = info.file_size / info.compress_size
        if ratio > ZIP_MAX_COMPRESSION_RATIO:
            raise ValueError(f"ZIP member exceeds safe compression ratio: {info.filename!r}")
    return member_path


def extract_zip(path: Path, settings: Settings, observer: RunObserver | None = None) -> ExtractionResult:
    started = time.perf_counter()
    output: list[str] = []
    methods: list[str] = []
    failures: list[str] = []
    supported_seen = 0
    total_uncompressed = 0

    try:
        archive = zipfile.ZipFile(path)
    except (OSError, zipfile.BadZipFile) as exc:
        raise ValueError(f"Invalid ZIP attachment: {exc}") from exc

    with archive, tempfile.TemporaryDirectory(prefix="idx-zip-") as temp_dir:
        infos = archive.infolist()
        if len(infos) > ZIP_MAX_MEMBERS:
            raise ValueError(f"ZIP archive has too many members: {len(infos)} > {ZIP_MAX_MEMBERS}")

        for index, info in enumerate(infos):
            if info.is_dir():
                continue
            member_path = _validate_zip_member(info)
            suffix = member_path.suffix.lower()
            if suffix not in ZIP_SUPPORTED_SUFFIXES:
                continue

            supported_seen += 1
            if supported_seen > ZIP_MAX_SUPPORTED_MEMBERS:
                raise ValueError(
                    f"ZIP archive has too many supported documents: {supported_seen} > {ZIP_MAX_SUPPORTED_MEMBERS}"
                )
            total_uncompressed += info.file_size
            if total_uncompressed > ZIP_MAX_TOTAL_UNCOMPRESSED_BYTES:
                raise ValueError("ZIP archive exceeds safe total uncompressed size limit")

            try:
                with archive.open(info, "r") as member_handle:
                    payload = member_handle.read(ZIP_MAX_MEMBER_UNCOMPRESSED_BYTES + 1)
                if len(payload) > ZIP_MAX_MEMBER_UNCOMPRESSED_BYTES:
                    raise ValueError("member exceeded safe read limit")

                safe_name = f"{index:03d}-{member_path.name}"
                member_local = Path(temp_dir) / safe_name
                member_local.write_bytes(payload)
                member_content_type = mimetypes.guess_type(member_path.name)[0] or "application/octet-stream"
                result = extract_document(member_local, member_content_type, settings, observer)
                if result.text:
                    output.append(f"===== ZIP MEMBER: {member_path.as_posix()} =====\n{result.text}")
                    methods.append(result.method)
            except Exception as exc:
                failures.append(f"{member_path.as_posix()}: {type(exc).__name__}: {exc}")

    if supported_seen == 0:
        raise ValueError("ZIP archive contains no supported document attachments")
    if not output:
        details = "; ".join(failures[:3])
        raise ValueError(f"ZIP archive contained supported documents but none could be extracted: {details}")

    if observer:
        observer.event(
            "extract",
            "ZIP extraction complete",
            file=path.name,
            supported_members=supported_seen,
            extracted_members=len(output),
            failed_members=len(failures),
            uncompressed_bytes=total_uncompressed,
            elapsed_seconds=f"{time.perf_counter() - started:.3f}",
        )

    method_set = list(dict.fromkeys(methods))
    method = "zip[" + "+".join(method_set) + "]"
    if failures:
        output.append(
            "===== ZIP EXTRACTION LIMITATIONS =====\n"
            + "\n".join(f"- {failure}" for failure in failures[:10])
        )
    return ExtractionResult(normalize_text("\n\n".join(output)), method)


def extract_document(
    path: Path,
    content_type: str,
    settings: Settings,
    observer: RunObserver | None = None,
) -> ExtractionResult:
    suffix = path.suffix.lower()
    ctype = (content_type or "").lower()
    if observer:
        observer.event(
            "extract",
            "selecting extractor",
            file=path.name,
            suffix=suffix,
            content_type=content_type,
            bytes=path.stat().st_size if path.exists() else None,
        )
    if suffix == ".zip" or ctype in {"application/zip", "application/x-zip-compressed"}:
        return extract_zip(path, settings, observer)
    if suffix == ".pdf" or "application/pdf" in ctype:
        return extract_pdf(path, settings, observer)
    if suffix in {".xlsx", ".xlsm"} or "spreadsheetml" in ctype:
        return extract_xlsx(path, settings, observer)
    if suffix == ".docx" or "wordprocessingml" in ctype:
        return extract_docx(path, observer)
    if suffix in {".html", ".htm"} or "text/html" in ctype:
        return extract_html(path, observer)
    if suffix in {".txt", ".csv", ".json", ".xml"} or ctype.startswith("text/"):
        return extract_text(path, observer)
    raise ValueError(f"Unsupported attachment type: suffix={suffix!r}, content_type={content_type!r}")